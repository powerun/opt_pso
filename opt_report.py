
import os
import time
import pandas as pd
import numpy as np
import datetime
import mysql_db
import zipfile
from docx import Document
from functools import reduce
from docx.shared import Inches
import matplotlib.pyplot as plt


plt.rcParams['font.sans-serif'] = ['SimHei']   # 用来正常显示中文标签
plt.rcParams['axes.unicode_minus'] = False   # 用来正常显示负号
document = Document('./templates/template_doc.docx')   # 获得模板


class OptReport(object):
    def __init__(self, prj):
        self.prj_list = prj
        self.logs = mysql_db.Logger('./logs/' + str(datetime.datetime.now().date()))
        self.my_sql = mysql_db.OptTemp(self.logs)
        self.my_sql.open_mysql('config.opt')

    @staticmethod
    def insert_df(df):
        cols_list = df.columns.tolist()
        table = document.add_table(rows=1, cols=df.shape[1])
        hdr_cells = table.rows[0].cells
        for j in range(df.shape[1]):
            hdr_cells[j].text = cols_list[j]
        for i in range(df.shape[0]):
            row_cells = table.add_row().cells
            for j in range(df.shape[1]):
                row_cells[j].text = str(df[cols_list[j]][i])

    @staticmethod
    def lists_combination(lists, code=''):
        def my_func(list1, list2):
            return [str(i) + code + str(j) for i in list1 for j in list2]
        return reduce(my_func, lists)

    def report_head(self):
        document.add_heading('机器配置测试报告', 0)
        document.add_paragraph('报告生成时间：%s' % time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()))
        document.add_heading('结果对比', level=1)
        rows = []
        for pr in self.prj_list:
            df = self.my_sql.query_data(''' SELECT t1.cold_need, t2.min_value, t2.peo_value 
                                            FROM time_need_tweb t1 JOIN opt_index t2 
                                            ON t1.ID='%s' AND t1.ID=t2.ID AND t1.hours=t2.hours 
                                            ORDER BY t1.hours ''' % pr)
            df = df[df['min_value'] > 0]
            opti_sum_kw = int(df['min_value'].sum())
            not_opti_sum_kw = int(df['peo_value'].sum())
            rows.append(['电量' + pr, not_opti_sum_kw, opti_sum_kw])
            df['cop'] = 3.517 * df['cold_need'] / df['min_value']
            df['cop2'] = 3.517 * df['cold_need'] / df['peo_value']
            rows.append(['EER' + pr, round(df['cop2'].mean(), 3), round(df['cop'].mean(), 3)])
        df = pd.DataFrame(rows, columns=['配置', '优化前', '优化后'])
        self.insert_df(df)

    def plot_state(self):
        for pr in self.prj_list:
            document.add_heading('机器类型信息和配置%s' % pr, level=1)
            df = self.my_sql.query_data(''' SELECT * FROM opt_plot WHERE ID='%s' ''' % pr)
            document.add_paragraph('tci=%s，teo=%s' % (df.at[0, 'tci'], df.at[0, 'teo']))
            for i, state in enumerate(df['state'].unique()):
                df_i = df[df['state'] == state]
                plt.plot(df_i['cold'], df_i['min_value'], lw=0.5, label='State_'+state)
            plt.legend(bbox_to_anchor=(1.005, 0), loc=3, borderaxespad=0, frameon=False)
            plt.title('开机状态最优运行功率对比')
            plt.xlabel('冷负荷需求（冷吨）')
            plt.ylabel('电功率kw')
            plt.tight_layout()
            plt.savefig('./templates/temp.png')
            plt.clf()
            document.add_picture('./templates/temp.png', width=Inches(6))

    def plot_cold(self):
        for i, pr in enumerate(self.prj_list):
            df = self.my_sql.query_data(''' SELECT t1.cold_need, t2.min_value, t2.peo_value 
                                            FROM (SELECT * FROM time_need_tweb WHERE ID='%s') t1 JOIN opt_index t2 
                                            ON t1.ID=t2.ID AND t1.hours=t2.hours 
                                            ORDER BY t1.hours ''' % pr)
            df['cop'] = 3.517*df['cold_need']/df['min_value']
            document.add_heading('冷量和cop散点图%s' % pr, level=1)
            plt.scatter(df['cold_need'], df['cop'], label='optimized')
            df['cop2'] = 3.517*df['cold_need']/df['peo_value']
            plt.scatter(df['cold_need'], df['cop2'], label='normal', marker='v')
            plt.legend()
            plt.xlabel('冷负荷需求（冷吨）')
            plt.ylabel('COP')  
            plt.tight_layout()          
            plt.savefig('./templates/temp.png')
            plt.clf()
            document.add_picture('./templates/temp.png', width=Inches(6))

    def plot_state2(self, report_folder):
        for pr in self.prj_list:
            if os.path.exists('opt_report/%s/%s' % (report_folder, pr)) is False:
                os.mkdir('opt_report/%s/%s' % (report_folder, pr))
            df = self.my_sql.query_data(''' SELECT t1.hours, t1.cold_need, IFNULL(t2.min_index,-1) as min_index, IFNULL(t2.peo_index,-1) as peo_index
                                            FROM (SELECT * FROM time_need_tweb WHERE ID='%s') t1
                                            LEFT JOIN opt_index t2
                                            ON t1.ID=t2.ID AND t1.hours=t2.hours
                                            ORDER BY t1.hours ''' % pr)
            df_types = self.my_sql.query_data(''' SELECT * FROM m_types WHERE ID='%s' ''' % pr)
            state_list = self.lists_combination([list(range(i + 1)) for i in df_types["m_nums"].tolist()])
            state_list = state_list[1:]
            for k in ['min_index', 'peo_index']:
                if 'min' in k:
                    document.add_heading('优化后各时段开机图%s' % pr, level=1)
                else:
                    document.add_heading('优化前各时段开机图%s' % pr, level=1)
                m = df[k].values
                t = [1] * (len(m))
                for j in range(len(m)):
                    i = m[j]
                    if i >= 0:
                        if j == 0:
                            plt.barh(i, t[j])
                        else:
                            plt.barh(i, t[j], left=(np.sum(t[:j])))
                plt.ylabel('机器状态')
                plt.xlabel('时间段')
                plt.yticks(np.arange(len(state_list)), labels=state_list)
                plt.tight_layout()
                plt.savefig('./templates/temp.png')
                plt.clf()
                document.add_picture('./templates/temp.png', width=Inches(6))

    def plot_picture(self, report_folder, period_hour=24):
        for pr in self.prj_list:
            df = self.my_sql.query_data(''' SELECT t1.hours, t1.cold_need, IFNULL(t2.min_index,-1) as min_index, IFNULL(t2.peo_index,-1) as peo_index
                                            FROM (SELECT * FROM time_need_tweb WHERE ID='%s') t1
                                            LEFT JOIN opt_index t2
                                            ON t1.ID=t2.ID AND t1.hours=t2.hours
                                            ORDER BY t1.hours ''' % pr)
            df_types = self.my_sql.query_data(''' SELECT * FROM m_types WHERE ID='%s' ''' % pr)
            state_list = self.lists_combination([list(range(i + 1)) for i in df_types["m_nums"].tolist()])
            state_list = state_list[1:]
            picture_date = datetime.datetime.strptime('2021-01-01', '%Y-%m-%d').date()
            for k in ['min_index', 'peo_index']:
                m_all = df[k].values
                for m_i in range(int(len(m_all)/period_hour)):
                    picture_name = picture_date + datetime.timedelta(days=m_i)
                    m = m_all[m_i*period_hour: (m_i+1)*period_hour]
                    if list(m).count(-1) < period_hour:
                        t = [1] * (len(m))
                        for j in range(len(m)):
                            i = m[j]
                            if i >= 0:
                                if j == 0:
                                    plt.barh(i, t[j], 0.3)
                                else:
                                    plt.barh(i, t[j], 0.3, left=(np.sum(t[:j])))
                        plt.ylabel('机器状态')
                        plt.xlabel('时间段H')
                        plt.yticks(np.arange(len(state_list)), labels=state_list)
                        plt.xticks(np.arange(period_hour), labels=[str(i) for i in range(period_hour)])
                        plt.tight_layout()
                        plt.savefig('opt_report/%s/%s/%s_%s.png' % (report_folder, pr, k[: 3], str(picture_name)[5:]))
                        plt.clf()

    @staticmethod
    def zip_result(report_path):
        if os.path.exists(report_path):
            z = zipfile.ZipFile(report_path + '.zip', 'w', zipfile.ZIP_DEFLATED)
            for dir_path, dir_names, file_names in os.walk(report_path):
                for filename in file_names:
                    path3 = dir_path.replace(report_path, '')
                    path3 = path3 and path3 + os.sep or ''
                    z.write(os.path.join(dir_path, filename), path3 + filename)
            z.close()

    def report_main(self, report_folder):
        report_path = os.path.join('opt_report', report_folder)
        if os.path.exists(report_path) is False:
            os.mkdir(report_path)
        self.report_head()
        self.plot_cold()
        self.plot_state()
        self.plot_state2(report_folder)
        document.save('%s/%s_%s.docx' % (report_path, '_'.join(self.prj_list), time.strftime("%Y-%m-%d-%H-%M-%S", time.localtime())))
        # self.plot_picture(report_folder)
        self.zip_result(report_path)


if __name__ == '__main__':
    OptReport(['3big_1small_y', 'p001']).report_main('2021-09-23 17-04-10')










